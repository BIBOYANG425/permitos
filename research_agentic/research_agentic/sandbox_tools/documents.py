"""read_pdf / read_docx / read_spreadsheet — run-workspace document readers.

Ported verbatim from the parent repo's documents.py (PR #38). Heavy parsers (PyMuPDF,
python-docx, openpyxl) are import-guarded — when absent they return a structured
'dependency_missing' result rather than raising. CSV needs no third-party dep. All paths
are guarded through _resolve_workspace_path and capped via _cap_text.
"""

from __future__ import annotations

import csv
from pathlib import Path
from typing import Any

from research_agentic.policy import (
    SandboxPolicy,
    _cap_text,
    _error,
    _exception_error,
    _resolve_workspace_path,
    _success,
)


def _path_or_error(policy: SandboxPolicy, path: str | Path) -> Path | dict[str, Any]:
    try:
        resolved = _resolve_workspace_path(policy, path)
    except TypeError as exc:
        return _error("error", "invalid_argument", str(exc), path=str(path))
    except ValueError as exc:
        return _error("error", "path_traversal", str(exc), path=str(path))
    if not resolved.exists():
        return _error("error", "file_not_found", "Document does not exist.", path=str(path))
    if not resolved.is_file():
        return _error("error", "not_a_file", "Document path is not a file.", path=str(path))
    return resolved


def read_pdf(policy: SandboxPolicy, path: str | Path) -> dict[str, Any]:
    checked = _path_or_error(policy, path)
    if isinstance(checked, dict):
        return checked
    try:
        import fitz
    except ImportError:
        return _error("unavailable", "dependency_missing", "PyMuPDF is not installed.", dependency="pymupdf")
    try:
        pages: list[dict[str, Any]] = []
        with fitz.open(checked) as document:
            for index, page in enumerate(document):
                pages.append({"page": index + 1, "text": page.get_text("text")})
        return _success("read", path=str(checked), page_count=len(pages),
                        text=_cap_text("\n".join(page["text"] for page in pages)), pages=pages)
    except Exception as exc:
        return _exception_error("pdf_read_failed", exc, path=str(checked))


def read_docx(policy: SandboxPolicy, path: str | Path) -> dict[str, Any]:
    checked = _path_or_error(policy, path)
    if isinstance(checked, dict):
        return checked
    try:
        import docx
    except ImportError:
        return _error("unavailable", "dependency_missing", "python-docx is not installed.", dependency="python-docx")
    try:
        document = docx.Document(str(checked))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        tables = [[[cell.text for cell in row.cells] for row in table.rows] for table in document.tables]
        return _success("read", path=str(checked), text=_cap_text("\n".join(paragraphs)),
                        paragraphs=paragraphs, tables=tables)
    except Exception as exc:
        return _exception_error("docx_read_failed", exc, path=str(checked))


def read_spreadsheet(policy: SandboxPolicy, path: str | Path) -> dict[str, Any]:
    checked = _path_or_error(policy, path)
    if isinstance(checked, dict):
        return checked
    suffix = checked.suffix.lower()
    if suffix == ".csv":
        return _read_csv(checked)
    if suffix in {".xlsx", ".xlsm"}:
        return _read_xlsx(checked)
    return _error("error", "unsupported_spreadsheet", "Unsupported spreadsheet format.", path=str(checked), suffix=suffix)


def _read_csv(path: Path) -> dict[str, Any]:
    try:
        with path.open(newline="") as handle:
            rows = list(csv.reader(handle))
        return _success("read", path=str(path), sheets=[{"name": path.stem, "rows": rows}], text=_cap_text(_rows_to_text(rows)))
    except Exception as exc:
        return _exception_error("csv_read_failed", exc, path=str(path))


def _read_xlsx(path: Path) -> dict[str, Any]:
    try:
        import openpyxl
    except ImportError:
        return _error("unavailable", "dependency_missing", "openpyxl is not installed.", dependency="openpyxl")
    try:
        workbook = openpyxl.load_workbook(path, data_only=True, read_only=True)
        sheets = []
        for sheet in workbook.worksheets:
            rows = [list(row) for row in sheet.iter_rows(values_only=True)]
            sheets.append({"name": sheet.title, "rows": rows})
        workbook.close()
        return _success("read", path=str(path), sheets=sheets,
                        text=_cap_text("\n\n".join(_rows_to_text(sheet["rows"]) for sheet in sheets)))
    except Exception as exc:
        return _exception_error("spreadsheet_read_failed", exc, path=str(path))


def _rows_to_text(rows: list[list[Any]]) -> str:
    return "\n".join("\t".join("" if cell is None else str(cell) for cell in row) for row in rows)
